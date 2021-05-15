from bs4 import BeautifulSoup
import re
import os
import tabula
import fitz
import docx
import csv

def get_tables(path):
    tabb = tabula.read_pdf(path, pages = "all", multiple_tables = True)
    name = path.strip('pdf')
    ccounter = 0
    for elem in tabb:
        ccounter += 1
        with open (name + str(ccounter) + '.xlsx', 'w', encoding = 'utf-8'):
            elem.to_excel(name + str(ccounter) + '.xlsx')
            elem.to_csv(name + str(ccounter) + '.csv')
    return ccounter
            
def get_images(path):
    doc = fitz.open(path)
    name = path.strip('pdf')
    ccounter = 0
    for page in doc:
        images = page.get_images()
        if images:
            for im in images:
                ccounter += 1
                pix = fitz.Pixmap(doc, im[0])
                pix.writeImage(name + str(ccounter) + '.png', output=None)
    return ccounter

def get_html(path, name):
    pathToScript = r'C:\Users\Anna\Python39\Scripts\pdf2txt.py'
    pathPDFinput = os.path.join(path)
    pathHTMLoutput = os.path.join(name)
    os.system(r'C:\Users\Anna\Python39\python.exe {} -o {} -S {} -t html'.format(pathToScript, pathHTMLoutput, pathPDFinput))
    
def num_search(thing): #для поиска тегов
    for e in thing.contents:
        if re.match('\(\d(\d)?\)', str(e)):
            return re.match('\((\d)+\)', str(e))
    return False

def n_search(t): #для содержания тегов
    if re.match(' *\(\d(\d)?\) *', str(t)):
        return re.match('\((\d)+\)', str(t))
    return False

def f_search(t): #для готового тега
    if 'style' in t.attrs:
        if re.match("font-family: CharterITC-(Bold)?Italic; .*", str(t['style'])):
            return re.match("font-family: CharterITC-(Bold)?Italic; .*", str(t['style']))
    return False

def get_italics(l, text):
    if f_search(l[0]):
        text.extend(l[0].contents)
        return get_italics(l[1:], text)
    else:
        return text
    
def p_search(list_of_soup):
    s = []
    for i in range (len(list_of_soup)):
        if num_search(list_of_soup[i]):
            if n_search(list_of_soup[i].contents[0]):
                tt = []
                sent = get_italics(list_of_soup[(i+1):], tt)
                s.append(sent)
    return s

def processing_s(s):
    ls = []
    coun = 0
    for elem in s:
        coun += 1
        sent = ''
        if elem != []:
            for ss in elem:
                if str(ss) != '<br/>':
                    if str(ss).endswith('-'):
                        sent += str(ss).strip('\n')
                    else:
                        sent += str(ss).strip('\n') + ' '
            ls.append('(' + str(coun) + ') ' + sent)
    return ls

def get_examples(path):
    n = path.strip('pdf') + 'html'
    get_html(path, n)
    with open(n, encoding = 'utf-8') as fp:
        soup = BeautifulSoup(fp, 'html.parser')
        span_list = soup.find_all('span')
        h = p_search(span_list)
        return (processing_s(h))
    
def write_examples(e, path):
    name = path.strip('pdf') + 'txt'
    if e!=[]:
        with open(name, 'w', encoding='utf-8') as f:
            for elem in e:
                f.write(elem + '\n\n')
    return len(e)

def make_docx(path, imn, tabn, exn):
    txtpath = path.strip('pdf') + 'txt'
    name = path.strip('pdf')
    doc = docx.Document()
    doc.add_heading('СТАТЬЯ: материалы', level=1)
    if exn != 0:
        doc.add_heading('Примеры', level=3)
        with open(txtpath, encoding='utf-8') as tet:
            t = tet.read()
            doc.add_paragraph(text=t)
        doc.add_page_break()
    if imn != 0:
        doc.add_heading('Картинки', level=3)
        for h in range(imn):
            doc.add_picture(name + str(h+1) + '.png')
            doc.add_page_break()
    if tabn != 0:
        doc.add_heading('Таблицы', level=3)
        for k in range(tabn):
            with open(name + str(k+1) + '.csv', encoding = 'utf-8') as csvfile:
                c = csv.DictReader(csvfile)
                headers = []
                for row in c:
                    colcounter = 0
                    for elem in row:
                        colcounter += 1
                table = doc.add_table(rows=1, cols=colcounter+2)
            with open(name + str(k+1) + '.csv', encoding = 'utf-8') as csvfile:
                c = csv.reader(csvfile)
                i = 0
                for row in c:
                    j = 0
                    for el in row:
                        if j == 0:
                            row_cells = table.add_row().cells
                        row_cells[j].text = el
                        j += 1
                    i += 1
            doc.add_page_break()
    doc.save(name + 'docx')
    
def main():
    p = input('insert the path')
    examples = get_examples(p)
    ex = write_examples(examples, p)
    imnum = get_images(p)
    tabl = get_tables(p)
    make_docx(p, imnum, tabl, ex)
    
        
main()
