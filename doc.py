import docx
import os
import csv
import tabula

def get_tables(path):
    tabb = tabula.read_pdf(path, pages = "all", multiple_tables = True)
    name = path.strip('pdf')
    ccounter = 0
    for elem in tabb:
        ccounter += 1
        with open (name + str(ccounter) + '.xlsx', 'w', encoding = 'utf-8'):
            elem.to_excel(name + str(ccounter) + '.xlsx')
            elem.to_csv(name + str(ccounter) + '.csv')
            
doc = docx.Document()
doc.add_heading('СТАТЬЯ')
with open('C:\Downloads\papka\chuikovaoiu-051.txt', encoding='utf-8') as tet:
    t = tet.read()
    doc.add_paragraph(text=t)
doc.add_page_break()
doc.add_picture('anastasyevdg-147.1.png')
doc.add_page_break()
doc.add_page_break()
with open('chuikovaoiu-051.14.csv', encoding = 'utf-8') as csvfile:
    c = csv.reader(csvfile)
    #nam = c.keys()
    table = doc.add_table(rows=1, cols=1)
    #hdr_cells = table.rows[0].cells
    #for i in range(len(nam)):
        #hdr_cells[i].text = nam[i]
    i = 0
    for row in c:
        col_cells = table.add_column(width=3).cells
        j = 0
        for el in row:
            if j == 0:
                #hdr_cells[i].text = el
                row_cells = table.add_row().cells
            #col_cells = add_column().cells
            row_cells[j].text = el
            j += 1
        i += 1

        
doc.save('test1.docx')
