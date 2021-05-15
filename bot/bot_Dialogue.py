from bs4 import BeautifulSoup
import re
import os
import tabula
import fitz
import docx
import json
import csv
import random
import pdfminer.high_level
from collections import Counter
import telebot
from telebot import types

bot = telebot.TeleBot('1752493898:AAGgYdZWuhtdQh3baJPCRl-nUXs72ZWp-kc')

#функции для обработки статьи-------------------------------------------------
def get_tables(path):
    tabb = tabula.read_pdf(path, pages = "all", multiple_tables = True)
    name = path.strip('pdf')
    ccounter = 0
    for elem in tabb:
        ccounter += 1
        with open (name + str(ccounter) + '.xlsx', 'w', encoding = 'utf-8'):
            elem.to_excel(name + str(ccounter) + '.xlsx')
            elem.to_csv(name + str(ccounter) + '.csv')
    return ccounter
            
def get_images(path):
    doc = fitz.open(path)
    name = path.strip('pdf')
    ccounter = 0
    for page in doc:
        images = page.get_images()
        if images:
            for im in images:
                ccounter += 1
                pix = fitz.Pixmap(doc, im[0])
                pix.writeImage(name + str(ccounter) + '.png', output=None)
    return ccounter

def get_html(path, name):
    pathToScript = r'C:\Users\kseni\AppData\Local\Programs\Python\Python38-3\
2\Scripts\pdf2txt.py'
    pathPDFinput = os.path.join(path)
    pathHTMLoutput = os.path.join(name)
    os.system(r'C:\Users\kseni\AppData\Local\Programs\Python\Python38-32\pyth\
on.exe {} -o {} -S {} -t html'.format(pathToScript, pathHTMLoutput, \
                                      pathPDFinput))
    
def num_search(thing): #для поиска тегов
    for e in thing.contents:
        if re.match('\(\d(\d)?\)', str(e)):
            return re.match('\((\d)+\)', str(e))
    return False

def n_search(t): #для содержания тегов
    if re.match(' *\(\d(\d)?\) *', str(t)):
        return re.match('\((\d)+\)', str(t))
    return False

def f_search(t): #для готового тега
    if 'style' in t.attrs:
        if re.match("font-family: CharterITC-(Bold)?Italic; .*",\
                    str(t['style'])):
            return re.match("font-family: CharterITC-(Bold)?Italic; .*", \
                            str(t['style']))
    return False

def get_italics(l, text):
    if f_search(l[0]):
        text.extend(l[0].contents)
        return get_italics(l[1:], text)
    else:
        return text
    
def p_search(list_of_soup):
    s = []
    for i in range (len(list_of_soup)):
        if num_search(list_of_soup[i]):
            if n_search(list_of_soup[i].contents[0]):
                tt = []
                sent = get_italics(list_of_soup[(i+1):], tt)
                s.append(sent)
    return s

def processing_s(s):
    ls = []
    coun = 0
    for elem in s:
        coun += 1
        sent = ''
        if elem != []:
            for ss in elem:
                if str(ss) != '<br/>':
                    if str(ss).endswith('-'):
                        sent += str(ss).strip('\n')
                    else:
                        sent += str(ss).strip('\n') + ' '
            ls.append('(' + str(coun) + ') ' + sent)
    return ls

def get_examples(path):
    n = path.strip('pdf') + 'html'
    get_html(path, n)
    with open(n, encoding = 'utf-8') as fp:
        soup = BeautifulSoup(fp, 'html.parser')
        span_list = soup.find_all('span')
        h = p_search(span_list)
        return (processing_s(h))
    
def write_examples(e, path):
    name = path.strip('pdf') + 'txt'
    if e!=[]:
        with open(name, 'w', encoding='utf-8') as f:
            for elem in e:
                f.write(elem + '\n\n')
    return len(e)

def make_docx(path, imn, tabn, exn):
    txtpath = path.strip('pdf') + 'txt'
    name = path.strip('pdf')
    doc = docx.Document()
    doc.add_heading('СТАТЬЯ: материалы', level=1)
    if exn != 0:
        doc.add_heading('Примеры', level=3)
        with open(txtpath, encoding='utf-8') as tet:
            t = tet.read()
            doc.add_paragraph(text=t)
        doc.add_page_break()
    if imn != 0:
        doc.add_heading('Картинки', level=3)
        for h in range(imn):
            doc.add_picture(name + str(h+1) + '.png')
            doc.add_page_break()
    if tabn != 0:
        doc.add_heading('Таблицы', level=3)
        for k in range(tabn):
            with open(name + str(k+1) + '.csv', encoding = 'utf-8') as \
                 csvfile:
                c = csv.DictReader(csvfile)
                headers = []
                for row in c:
                    rowcounter = 0
                    for elem in row:
                        rowcounter += 1
                table = doc.add_table(rows=1, cols=rowcounter)
            with open(name + str(k+1) + '.csv', encoding = 'utf-8') as \
                 csvfile:
                c = csv.reader(csvfile)
                i = 0
                for row in c:
                    j = 0
                    for el in row:
                        if j == 0:
                            row_cells = table.add_row().cells
                        row_cells[j].text = el
                        j += 1
                    i += 1
            doc.add_page_break()
    doc.save(name + 'docx')
    
def main(p):
    examples = get_examples(p)
    ex = write_examples(examples, p)
    imnum = get_images(p)
    tabl = get_tables(p)
    make_docx(p, imnum, tabl, ex)
    p = p.rstrip("pdf")
    p += "docx"
    return p

#функции для других двух разделов бота----------------------------------------
def obr_keyword(word):
    w = word.lower()
    w = w.replace("keyword:","", 1)
    w = w.strip(' ')
    final = "Со словом \'" + w + "\' следующие статьи: \n \n"
    anti_final = "Простите, у нас нет такого слова... \nПопробуйте проверить,\
 вдруг в вашем сообщении есть опечатки"

    with open('keywords.json', 'r', encoding='utf-8') as f:
        text = json.load(f)
        if w in text:
            a = text[w]
            for i in a:
                final += "🌲" + i[1]
                final += ", "
                final += i[0]
                final += " \n"
            return  final
        else:
            return anti_final


def obr_articles(tex):
    links = csv.DictReader(open('links_to_articles.csv', encoding='utf-8'))
    links_list = []
    title = tex.lower()
    title = title.replace("статья:","", 1)
    title = title.strip(' ')
    var_final = 0
    for row in links: 
        new = str(row['title']) 
        new = new.lower() 
        if title in new: 
            links_list.append(str(row['link'])) 
    line = len(links_list)
    if line == 0:
        links_list = ["0"]
    return links_list, line

def randomizer():
    with open('authors_and_titles.json', 'r', encoding='utf-8') as f:
        text = json.load(f)
        a = random.randrange(1, 326)
        a = str(a)
        return text[a]

def randomizer_words():
    with open('random_text.json', 'r', encoding='utf-8') as f:
        text = json.load(f)
        a = random.randrange(1, 12)
        a = str(a)
        return text[a]

#команды бота-----------------------------------------------------------------
@bot.message_handler(commands = ['start'])
def starter(message):
    bot.send_message(message.from_user.id, "Наш бот приветствует вас! Для того\
, чтобы попасть в Стартовое меню, нажмите /menu. \nТам вы можете увидеть кноп\
очки. Пройдёмся по каждой из них. \n \n“*Обработка статьи* 🍏”. Здесь вы можете п\
олучить либо список примеров, картинок и таблиц из этой статьи, либо топ-25 са\
мых её частотных слов. \n \n“*Достать статью Диалога* 🥑”. Эта функция позволяет \
“пробивать” статьи Диалога за 2020 год по нашей базе. Вы получите ссылку на \
гугл папки с материалами по данной статье. \n \n“*Интересные факты* 🥝”. Здесь ле\
жит статистика по статьям Диалога за последние 5 лет. Можно посмотреть следующ\
ее: список из 5 самых пишущих авторов, список из 20 самых популярных ключевых \
слов. Также можно найти все статьи с определённым ключевым словом; а если вы н\
е знаете, какую статью почитать, то введите Random, и наш рандомайзер сгенерир\
ует вам случайную статью! \n \nНам кажется, вы с легкостью разберетесь в навиг\
ации по меню нашего бота :) Но всё же, если вы запутались в функциях бота и хо\
тите получить более подробную информацию о них, нажмите /help. \n \nДобро пожа\
ловать в мир лингвистический статей! А теперь нажмите /menu :)",
                     parse_mode= 'Markdown')
    chat_id = message.chat.id

@bot.message_handler(commands = ['help'])
def starter(message):
    bot.send_message(message.from_user.id, "В Стартовом меню (которое вы может\
е увидеть, нажав /menu) вы можете увидеть кнопочки. Пройдёмся по каждой из них\
. \n \n“*Обработка статьи* 🍏”. Здесь находятся две функции - “хэндаут” и “час\
тотные слова”. Нажимая на каждую из них, вы можете отправлять статьи (пожалуй\
ста, присылайте их в формате pdf!) и получать либо список примеров, картинок и\
 таблиц из этой статьи, либо топ-25 самых её частотных слов. \n \n“*Достать стать\
ю Диалога* 🥑”. Эта функция позволяет “пробивать” статьи Диалога за 2020 год\
 по нашей базе. Введите название статьи или его часть в формате статья: (назва\
ние), и вы получите ссылки на гугл папки со всеми статьями, где есть это назва\
ние или слова из названия. А вот что лежит в папках - написано в папках :) \n \n\
“*Интересные факты* 🥝”. Здесь лежит статистика по статьям Диалога за последни\
е 5 лет (2016-2020 годы). Можно посмотреть следующее: список из 5 самых пишущи\
х авторов, список из 20 самых популярных ключевых слов. Также можно найти все \
статьи с определённым ключевым словом (введите keyword: (ключевое слово)); а е\
сли вы не знаете, какую статью почитать, то введите Random, и наш рандомайзер \
сгенерирует вам случайную статью! \n \nНебольшая подсказка - для нескольких фу\
нкций было бы удобнее просто вводить команду, а не проходить путь с кнопочками\
 со Стартового меню: \n🌴написав keyword: (ключевое слово), вы получите список\
 всех статей с этим ключевым словом \n🌴написав статья: (название статьи), вы \
получите ссылку на гугл-папку с некоторыми файлами по этой статье \n🌴написав \
Random, вы получите случайную статью \n🌴для всего остального лучше использова\
ть Стартовое меню",parse_mode= 'Markdown')
    chat_id = message.chat.id

@bot.message_handler(commands=['menu'])
def menu(message):
    start_menu = types.ReplyKeyboardMarkup(True, True)
    start_menu.row('Обработка статьи 🍏')
    start_menu.row('Достать статью Диалога 🥑')
    start_menu.row('Интересные факты 🥝')
    bot.send_message(message.chat.id, 'Вы попали в Стартовое меню',\
reply_markup=start_menu)

#текстовые взаимодействия бота------------------------------------------------
@bot.message_handler(content_types=['text'])
def handle_text(message):
    if message.text == 'Обработка статьи 🍏':
        source_language_menu = types.ReplyKeyboardMarkup(True, True)
        source_language_menu.row('Частотные слова', 'Хэндаут')
        source_language_menu.row('Назад')
        bot.send_message(message.chat.id,"*Частотные слова* - тут вы получите \
25 самых частотных слов вашей статьи \n*Хэндаут* - а тут вы получите файл с “х\
эндаутом” вашей статьи (вынесенными из статьи картинками, примерами и таблицам\
и)",parse_mode= 'Markdown', reply_markup=source_language_menu)

    elif message.text == 'Достать статью Диалога 🥑':
        B_menu = types.ReplyKeyboardMarkup(True, True)
        B_menu.row('Назад')
        bot.send_message(message.chat.id, 'Напишите, пожалуйста, название стат\
ьи вот в таком формате \n  \nстатья: ____________', reply_markup=B_menu)
        
    elif "статья:" in message.text or "Статья:" in message.text:
        k = message.text
        answer = obr_articles(message.text)[0]
        line = obr_articles(message.text)[1]
        if line == 0:
            bot.send_message(message.chat.id, "Простите, у нас нет такой стать\
и... \nПопробуйте проверить, вдруг в вашем сообщении есть опечатки")
        elif line > 1:
            mess = "Вау, у нас целых несколько статей с таким названием!  \nВо\
т они все: \n"
            for i in range(len(answer)):
                if i == 0:
                    mess += answer[i]
                    mess +=  " \n"
                if i != 0:
                    mess += "🌿"
                    mess +=  " \n"
                    mess += answer[i]
                    mess +=  " \n"
            bot.send_message(message.chat.id, mess)
        else:
            markup = types.InlineKeyboardMarkup()
            btn_article= types.InlineKeyboardButton(text='🥑🥑🥑', url\
                                                    = answer[0])
            markup.add(btn_article)
            bot.send_message(message.chat.id, "Ура, тыкните на кнопку, чтобы \
забрать материалы", reply_markup = markup)
            bot.send_message(message.chat.id, "Теперь вы можете попросить инф\
ормацию по другим статьям или вернуться назад в меню")

    elif message.text == 'Интересные факты 🥝':
        C_menu = types.ReplyKeyboardMarkup(True, True)
        C_menu.row('Авторы', 'Ключевые слова', 'Поиск')
        C_menu.row('Random', 'Назад')
        bot.send_message(message.chat.id, "Здесь собрана интересная информация\
 по статьям Диалога за последние 5 лет: \n \n🥝Чтобы получить список из 5 самы\
х пишущих авторов: \nнажмите *Авторы* \n🥝Чтобы получить топ-20 keywords: \nна\
жмите *Ключевые слова* \n🥝Чтобы найти все статьи с конкретным keyword: \nнажм\
ите *Поиск* \n \nА если вы вдруг не знаете, какую бы статью почитать - предлаг\
аем нажать *Random*", parse_mode= 'Markdown', reply_markup=C_menu)

    elif message.text == 'Авторы':
        B_menu = types.ReplyKeyboardMarkup(True, True)
        B_menu.row('Назад!')
        bot.send_message(message.chat.id, 'Количество кактусов соответствует к\
оличеству написанных статей (за последние 5 лет): \n \nАпресян В. Ю.:   🌵🌵\
🌵🌵🌵🌵🌵🌵 \nИнькова О. Ю.:    🌵🌵🌵🌵🌵🌵 \nЛевонтина И. Б.: 🌵🌵🌵🌵🌵🌵 \
\nКустова Г. И.:        🌵🌵🌵🌵🌵 \nЛукашевич Н.В.:  🌵🌵🌵🌵🌵',\
 reply_markup=B_menu)

    elif message.text == 'Ключевые слова':
        B_menu = types.ReplyKeyboardMarkup(True, True)
        B_menu.row('Назад!')
        bot.send_message(message.chat.id, '*Топ-20 самых популярных keywords:*\
\n_Russian, Russian language_ \n \nCorpus linguistics, semantics, parallel cor\
pora, word embeddings, natural language processing, sentiment analysis, polyse\
my, word2vec, distributional semantics, BERT, prosody, machine learning, lemma\
tization, corpus, thesaurus, phonetics, spoken language, semantic similarity, \
multimodal communication, named entity recognition.',parse_mode= 'Markdown',\
                         reply_markup=B_menu)

    elif message.text == 'Поиск':
        B_menu = types.ReplyKeyboardMarkup(True, True)
        B_menu.row('Назад!')
        bot.send_message(message.chat.id, 'Введите ключевое слово (на английс\
ком) в следующем формате: \n \nkeyword: _________', reply_markup=B_menu)

    elif "keyword" in message.text or "Keyword" in message.text:
        B_menu = types.ReplyKeyboardMarkup(True, True)
        B_menu.row('Назад!')
        bot.send_message(message.chat.id, obr_keyword(message.text), \
                         reply_markup=B_menu)
        
    elif message.text == 'Random':
        L_menu = types.ReplyKeyboardMarkup(True, True)
        L_menu.row('Назад!', 'Хочу ещё..')
        a = randomizer()
        b = randomizer_words()
        random_message = b + " \n" + a[1] + ", " + a[0]
        bot.send_message(message.chat.id, random_message, reply_markup=L_menu)

    elif message.text == 'Хочу ещё..':
        L_menu = types.ReplyKeyboardMarkup(True, True)
        L_menu.row('Назад!', 'Хочу ещё..')
        a = randomizer()
        b = randomizer_words()
        random_message = b + " \n" + a[1] + ", " + a[0]
        bot.send_message(message.chat.id, random_message, reply_markup=L_menu)

    elif message.text == 'Назад!':
        C_menu = types.ReplyKeyboardMarkup(True, True)
        C_menu.row('Авторы', 'Ключевые слова', 'Поиск')
        C_menu.row('Random', 'Назад') 
        bot.send_message(message.chat.id, 'Вы вернулись в меню интересных фак\
тов ', reply_markup=C_menu)

    elif message.text == 'Частотные слова':
        source_language_menu = types.ReplyKeyboardMarkup(True, True)
        source_language_menu.row('Частотные слова', 'Хэндаут')
        source_language_menu.row('Назад')
        bot.send_message(message.chat.id, 'Отправьте файл в формате pdf с под\
писью “частотность”', reply_markup=source_language_menu)

    elif message.text == 'Хэндаут':
        source_language_menu = types.ReplyKeyboardMarkup(True, True)
        source_language_menu.row('Частотные слова', 'Хэндаут')
        source_language_menu.row('Назад')
        bot.send_message(message.chat.id, 'Отправьте файл в формате pdf с под\
писью “хэндаут”', reply_markup=source_language_menu)

    elif message.text == 'Назад':
        menu(message)

    else:
        bot.send_message(message.chat.id, 'Прости, я не понимаю, попробуй нач\
ать с /menu')

#взаимодействие бота с отправленным документом--------------------------------
@bot.message_handler(content_types=['document'])
def handle_docs(message):
    if message.caption == "частотность":
        try:
            file_info = bot.get_file(message.document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            src = 'C:/Downloads/papka/' + message.document.file_name;
            with open(src, 'wb') as new_file:
                new_file.write(downloaded_file)

            bot.reply_to(message, "Секундочку, сейчас мы это обработаем")
            n = 25
            pdfdir = src

            textpdf = pdfminer.high_level.extract_text(pdfdir)
            textpdf = textpdf.replace('\n', ' ')
# все абзацы превратили в пробелы, текст теперь - сплошная строка
# очистка от переносов
            if '- ' in textpdf: 
                j = 0
                k = len(textpdf)
                while j < k-2:
                    if textpdf[j] == '-' and textpdf[j+1] == ' ': 
                        textpdf = textpdf[:j] + textpdf[j+2:] 
                        k -= 2 
                    else:
                        j += 1 

            with open('stop_all.txt', encoding='utf-8') as f:
                stop_list = f.read() 
            stop_words = []
            for word in stop_list.split():
                word = word.strip('!~./\\"\'*(),;:?%^<>&-=+\n»«')
                stop_words.append(word) 

            article = textpdf
            words = []
            for word in article.split():
                word = word.strip('!~./\\"\'*(),;:?%^<>&-=+\n»«')
                if word.lower() not in stop_words and word != '':
                    words.append(word)

            mostnwithfreq = [] 
            mostn = [] 
            mostnwithfreq = Counter(words).most_common(n)
            for t in mostnwithfreq:
                mostn.append(t[0])
            mostn = str(mostn)
            mostn = mostn.strip("[]")
            final_line = "Вот и 25 самых частотных слов этого файла: \n \n" + \
                         mostn
            bot.send_message(message.chat.id, final_line)
        except Exception as e:
            bot.reply_to(message, e)


    if message.caption == "хэндаут":
        try:
            file_info = bot.get_file(message.document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            src = 'C:/Downloads/papka/' + message.document.file_name;
            with open(src, 'wb') as new_file:
                new_file.write(downloaded_file)
            
            bot.reply_to(message, "Секундочку, сейчас мы это обработаем")

            f = open(main(src),"rb")
            bot.send_document(message.chat.id,f)
            bot.send_message(message.chat.id, 'Ура, заберите ваш хэндаут!!')
            
        except Exception as e:
            bot.reply_to(message, e)

#чтобы бот реагировал на сообщения--------------------------------------------
bot.polling(none_stop=True, interval=0) 




